"""
DOCX Generator Utility
=======================
Converts a legislative analysis memo into a formal Word memorandum.

Document structure:
  - MEMORANDUM title (Senate red, centred)
  - TO / FROM / DATE / RE header block (two-column table, no borders)
  - Senate red horizontal rule
  - Body: headings, bold inline text, bullet lists
  - Confidentiality footer

Markdown conventions handled:
  ## Heading      → Word Heading 2
  # Heading       → Word Heading 1
  1. SECTION      → Word Heading 2
  ALL CAPS LINE   → Word Heading 2
  - bullet / * bullet  → Word List Bullet
  **bold text**   → bold run within a paragraph
  ---             → skipped (markdown rule)
"""

import re
from io import BytesIO
from datetime import datetime


def generate_docx_buffer(
    memo_text:       str,
    bill_id:         str,
    session:         str  = "",
    testifier_count: int  = 0,
    memo_style:      str  = "full_analysis",
) -> BytesIO:
    """
    Build a formal Word .docx memorandum and return a seeked BytesIO buffer.
    Raises ImportError if python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    # ── Colour palette ────────────────────────────────────────────────────────
    SENATE_RED  = RGBColor(0x7A, 0x1A, 0x1A)
    DARK_RED    = RGBColor(0x5A, 0x00, 0x00)
    DARK_GREY   = RGBColor(0x37, 0x41, 0x51)
    MID_GREY    = RGBColor(0x6B, 0x72, 0x80)
    LIGHT_GREY  = RGBColor(0x9C, 0xA3, 0xAF)

    STYLE_LABELS = {
        "full_analysis":     "Full Analysis",
        "quick_summary":     "Quick Summary",
        "amendments_focus":  "Amendments Focus",
        "opposition_report": "Opposition Report",
    }

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)

    # ── Helper: add a horizontal rule via paragraph bottom border ─────────────
    def add_hr(para, color_hex="7A1A1A", sz="8"):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    sz)
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Helper: remove all borders from a table cell ─────────────────────────
    def clear_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        bdr  = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            bdr.append(el)
        tcPr.append(bdr)

    # ── Helper: render **bold** markers as bold runs ──────────────────────────
    BOLD_RE   = re.compile(r'\*\*(.+?)\*\*')
    BULLET_RE = re.compile(r'^[-*•]\s+(.+)')
    HEADING_RE = re.compile(
        r'^(#{1,3}\s+.+|\d+\.\s+[A-Z][A-Z\s/()\-]{3,}|[A-Z][A-Z\s/()\-]{4,}[A-Z)])$'
    )

    def add_rich_text(para, text: str, size_pt=11, default_colour=None):
        """Append text to a paragraph, rendering **bold** spans as bold runs."""
        parts = BOLD_RE.split(text)
        for j, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            run.font.size = Pt(size_pt)
            if j % 2 == 1:          # captured bold group
                run.bold = True
            if default_colour:
                run.font.color.rgb = default_colour

    # ── MEMORANDUM title ──────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(6)
    t_run = title_p.add_run("MEMORANDUM")
    t_run.bold           = True
    t_run.font.size      = Pt(20)
    t_run.font.color.rgb = SENATE_RED

    # ── TO / FROM / DATE / RE header (two-column borderless table) ────────────
    re_value = bill_id
    if session:
        re_value += f"  ·  Session: {session}"
    re_value += f"  ·  {STYLE_LABELS.get(memo_style, memo_style)}"
    if testifier_count:
        re_value += f"  ·  {testifier_count} Testifier(s)"

    header_rows = [
        ("TO:",    "Members, Maryland Senate"),
        ("FROM:",  "Senate Legislative Analysis Office"),
        ("DATE:",  datetime.now().strftime("%B %d, %Y")),
        ("RE:",    re_value),
    ]

    tbl = doc.add_table(rows=len(header_rows), cols=2)
    tbl.autofit = False

    # Set column widths: label col narrow, value col wide
    col_widths = [Inches(0.9), Inches(5.15)]
    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
            clear_cell_borders(cell)

    for i, (label, value) in enumerate(header_rows):
        lbl_cell, val_cell = tbl.rows[i].cells

        # Label (right-aligned, dark red, bold)
        lbl_p = lbl_cell.paragraphs[0]
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lbl_p.paragraph_format.space_before = Pt(2)
        lbl_p.paragraph_format.space_after  = Pt(2)
        lbl_run = lbl_p.add_run(label)
        lbl_run.bold           = True
        lbl_run.font.size      = Pt(11)
        lbl_run.font.color.rgb = DARK_RED

        # Value
        val_p = val_cell.paragraphs[0]
        val_p.paragraph_format.space_before = Pt(2)
        val_p.paragraph_format.space_after  = Pt(2)
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(11)
        if i == 3:  # RE: stands out
            val_run.bold           = True
            val_run.font.color.rgb = DARK_GREY

    # ── Horizontal rule ───────────────────────────────────────────────────────
    spacer1 = doc.add_paragraph()
    spacer1.paragraph_format.space_before = Pt(6)
    spacer1.paragraph_format.space_after  = Pt(0)
    add_hr(spacer1, "7A1A1A", "12")

    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_before = Pt(0)
    spacer2.paragraph_format.space_after  = Pt(8)

    # ── Body ─────────────────────────────────────────────────────────────────
    lines = memo_text.split("\n")
    i     = 0
    while i < len(lines):
        raw      = lines[i]
        stripped = raw.strip()
        i       += 1

        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(3)
            continue

        # Skip markdown / memo horizontal rules
        if stripped.startswith("---") or stripped.startswith("==="):
            continue

        # Headings
        if HEADING_RE.match(stripped):
            clean = re.sub(r'^[#\d.\s]+', '', stripped).strip() or stripped
            h = doc.add_heading(clean, level=2)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after  = Pt(5)
            for run in h.runs:
                run.font.color.rgb = DARK_RED
                run.font.size      = Pt(12)
                run.bold           = True
            continue

        # Bullet points
        bm = BULLET_RE.match(stripped)
        if bm:
            try:
                p = doc.add_paragraph(style="List Bullet")
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
            add_rich_text(p, bm.group(1), size_pt=11)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            continue

        # Regular body paragraph
        p = doc.add_paragraph()
        add_rich_text(p, stripped, size_pt=11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_hr = doc.add_paragraph()
    add_hr(footer_hr, "BFBFBF", "4")

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(4)
    f_run = footer_p.add_run(
        "Maryland General Assembly  ·  Senate Legislative Analysis  ·  CONFIDENTIAL — Internal Use Only"
    )
    f_run.font.size      = Pt(8)
    f_run.font.color.rgb = LIGHT_GREY
    f_run.italic         = True

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
